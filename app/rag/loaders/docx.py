"""
DOCX document loader using python-docx.
"""

import hashlib
from pathlib import Path

from docx import Document as DocxDocument

from app.core.errors import DocumentProcessingError
from app.models.document import Document, DocumentMetadata

from .base import BaseDocumentLoader


class DOCXLoader(BaseDocumentLoader):
    """DOCX document loader using python-docx."""

    async def load(self) -> Document:
        """
        Load and parse DOCX document.

        Returns:
            Document: Parsed DOCX document.

        Raises:
            DocumentProcessingError: If DOCX loading fails.
        """
        self._validate_file_exists()
        self._validate_file_readable()

        try:
            # Read file content for hash
            content_bytes = self.file_path.read_bytes()
            file_hash = hashlib.sha256(content_bytes).hexdigest()
            file_size = len(content_bytes)

            # Open DOCX with python-docx
            docx_document = DocxDocument(self.file_path)

            # Extract text from paragraphs
            paragraphs = [para.text for para in docx_document.paragraphs if para.text.strip()]

            # Combine all paragraphs
            full_text = "\n\n".join(paragraphs)

            # Create metadata
            metadata = DocumentMetadata(
                filename=self.file_path.name,
                file_size_bytes=file_size,
                file_hash=file_hash,
                source_type="docx",
            )

            # Create document
            document = Document(content=full_text.strip(), metadata=metadata)

            return document

        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to load DOCX file: {self.file_path.name}",
                details={"filename": self.file_path.name, "error": str(e)},
            )
